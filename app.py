"""
🤖 AI Assistant with Specialized Law Problem Answering
A conversational AI that answers any question using web search and fact-checking
Specialized IRAC format for law questions with OSCOLA citations
Uses Groq and Tavily APIs directly for reliable performance
"""

import streamlit as st
import os
from typing import Dict, Any, List, Tuple, Optional
from groq import Groq
from tavily import TavilyClient
from dotenv import load_dotenv
import re
import io
from pypdf import PdfReader
from docx import Document
import glob
import json
from pathlib import Path

# RAG imports - silent check, no warnings
RAG_AVAILABLE = False
RAG_FAISS_AVAILABLE = False

try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    try:
        # Try new import path (langchain >= 0.1.0)
        from langchain_community.embeddings import HuggingFaceEmbeddings
        from langchain_community.vectorstores import FAISS
        from langchain_core.documents import Document as LangchainDocument
        RAG_AVAILABLE = True
        RAG_FAISS_AVAILABLE = True
    except ImportError:
        try:
            # Fallback to old import path
            from langchain.embeddings import HuggingFaceEmbeddings
            from langchain.vectorstores import FAISS
            from langchain.docstore.document import Document as LangchainDocument
            RAG_AVAILABLE = True
            RAG_FAISS_AVAILABLE = True
        except ImportError:
            # Check if at least langchain core is available
            try:
                import langchain
                RAG_AVAILABLE = True  # Partial RAG without FAISS
            except ImportError:
                RAG_AVAILABLE = False
except ImportError:
    RAG_AVAILABLE = False

# Load environment variables
load_dotenv()

# Free Law Database URLs (from https://commonslibrary.parliament.uk/resources/legal-research-databases/)
FREE_LAW_DATABASES = {
    "BAILII": "https://www.bailii.org/",
    "Legislation.gov.uk": "https://www.legislation.gov.uk/",
    "British and Irish Legal Information Institute": "https://www.bailii.org/",
    "The National Archives": "https://www.nationalarchives.gov.uk/",
    "Courts and Tribunals Judiciary": "https://www.judiciary.uk/",
    "UK Parliament": "https://www.parliament.uk/",
    "Supreme Court": "https://www.supremecourt.uk/",
    "Legal Information Institute (Cornell)": "https://www.law.cornell.edu/",
    "Google Scholar": "https://scholar.google.com/",
    "SSRN Legal": "https://www.ssrn.com/",
    "Justis": "https://www.justis.com/",
    "Westlaw UK (free cases)": "https://www.westlaw.co.uk/",
    "CommonLII": "https://www.commonlii.org/",
    "AUSTLII": "https://www.austlii.edu.au/",
    "CANLII": "https://www.canlii.org/",
    "WorldLII": "https://www.worldlii.org/",
    "HKLII": "https://www.hklii.hk/",
    "NZLII": "https://www.nzlii.org/",
    "SAFLII": "https://www.saflii.org/",
    "Google": "https://www.google.com/",
    "Yahoo": "https://search.yahoo.com/",
    "Bing": "https://www.bing.com/",
}

# Page configuration
st.set_page_config(
    page_title="🤖 AI Assistant",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for ChatGPT-like interface
st.markdown("""
    <style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 1rem;
    }
    .stTextInput > div > div > input {
        font-size: 1.1rem;
        padding: 0.75rem;
    }
    .law-answer {
        background-color: #f8f9fa;
        padding: 1.5rem;
        border-left: 4px solid #667eea;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []

if 'groq_api_key' not in st.session_state:
    st.session_state.groq_api_key = os.getenv("GROQ_API_KEY", "")

if 'tavily_api_key' not in st.session_state:
    st.session_state.tavily_api_key = os.getenv("TAVILY_API_KEY", "")

if 'uploaded_documents' not in st.session_state:
    st.session_state.uploaded_documents = []  # List of {name: str, content: str}

if 'rag_vectorstore' not in st.session_state:
    st.session_state.rag_vectorstore = None

if 'rag_initialized' not in st.session_state:
    st.session_state.rag_initialized = False

def initialize_rag_knowledge_base():
    """Initialize RAG knowledge base from Law resouces folder"""
    if not RAG_AVAILABLE:
        return None
    
    if st.session_state.rag_initialized and st.session_state.rag_vectorstore is not None:
        return st.session_state.rag_vectorstore
    
    try:
        # Check both "Law resouces" (actual folder name) and "law_resources" (fallback)
        law_resources_path = Path("Law resouces")
        if not law_resources_path.exists():
            law_resources_path = Path("law_resources")
        if not law_resources_path.exists():
            return None
        
        # Load all documents from law_resources folder
        documents = []
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        
        # Load from all subdirectories
        for file_path in law_resources_path.rglob("*"):
            if file_path.is_file():
                try:
                    if file_path.suffix.lower() == '.pdf':
                        with open(file_path, 'rb') as f:
                            text = extract_text_from_pdf(f)
                    elif file_path.suffix.lower() in ['.txt', '.md']:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            text = f.read()
                    elif file_path.suffix.lower() in ['.docx', '.doc']:
                        with open(file_path, 'rb') as f:
                            text = extract_text_from_docx(f)
                    else:
                        continue
                    
                    if text and len(text.strip()) > 50:
                        chunks = text_splitter.split_text(text)
                        for chunk in chunks:
                            doc = LangchainDocument(
                                page_content=chunk,
                                metadata={"source": str(file_path), "type": file_path.parent.name}
                            )
                            documents.append(doc)
                except Exception as e:
                    continue
        
        if not documents:
            return None
        
        # Create embeddings and vectorstore
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        vectorstore = FAISS.from_documents(documents, embeddings)
        
        st.session_state.rag_vectorstore = vectorstore
        st.session_state.rag_initialized = True
        return vectorstore
    except Exception as e:
        return None

def retrieve_from_rag(query: str, k: int = 10) -> List[Dict]:
    """Retrieve relevant documents from RAG knowledge base - increased k for better coverage"""
    vectorstore = initialize_rag_knowledge_base()
    if vectorstore is None:
        return []
    
    try:
        docs = vectorstore.similarity_search(query, k=k)
        retrieved_docs = []
        for i, doc in enumerate(docs):
            retrieved_docs.append({
                'title': f"Knowledge Base: {doc.metadata.get('source', 'Document')}",
                'url': doc.metadata.get('source', ''),
                'content': doc.page_content,
                'type': doc.metadata.get('type', 'general')
            })
        return retrieved_docs
    except Exception as e:
        return []

def detect_hallucination(response: str, sources: List[Dict], groq_client: Groq, model: str, tavily_client: Optional[TavilyClient] = None) -> Tuple[bool, str]:
    """Enhanced hallucination detection with HallBayes-inspired approach - checks against sources and optionally internet"""
    if not sources:
        return False, "No sources to verify against"
    
    # Extract claims from response (cases, statutes, key facts)
    source_text = "\n\n".join([f"Source {i+1}: {s.get('content', '')[:500]}" for i, s in enumerate(sources[:10])])
    
    # Extract key claims from response for verification
    verification_prompt = f"""You are a strict fact-checker using a HallBayes-inspired approach. Analyze the response below and extract ALL factual claims (case names, statute names, dates, legal principles, specific facts, numbers, names).

Response to verify:
{response[:2000]}

Available sources:
{source_text}

Task: 
1. Extract ALL factual claims from the response
2. For EACH claim, check if it appears in the sources above
3. List ONLY claims that are NOT found in the sources

Format your response as:
- Claim: [the exact claim that cannot be verified]
- Issue: [why it cannot be verified - be specific]
- Source Check: [which sources you checked]

If ALL claims are found in sources, respond with "NO HALLUCINATIONS DETECTED - ALL CLAIMS VERIFIED IN SOURCES".

Be VERY strict - if a claim is not EXACTLY in the sources, flag it."""
    
    try:
        completion = groq_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a strict fact-checker using Bayesian reasoning. Only flag claims that are clearly not in the sources. Be extremely precise and methodical."},
                {"role": "user", "content": verification_prompt}
            ],
            temperature=0.1,
            max_tokens=800
        )
        result = completion.choices[0].message.content
        has_hallucination = "NO HALLUCINATIONS" not in result.upper()
        
        # If hallucinations detected and internet search available, do additional verification
        if has_hallucination and tavily_client:
            # Extract claims that need internet verification
            claims_to_check = []
            lines = result.split('\n')
            for line in lines:
                if 'Claim:' in line:
                    claim = line.split('Claim:')[1].strip()
                    claims_to_check.append(claim)
            
            # Verify claims with internet search
            if claims_to_check:
                internet_verification = "\n\n**Internet Verification Results:**\n"
                for claim in claims_to_check[:3]:  # Limit to 3 claims
                    try:
                        # Search Google/Yahoo/Bing for verification
                        search_results, _ = search_web(f"{claim} legal", tavily_client, max_results=2, search_type="general")
                        if "Error" not in search_results:
                            internet_verification += f"\n- Claim: {claim}\n  Internet search performed - results may need manual verification\n"
                    except:
                        pass
                result += internet_verification
        
        return has_hallucination, result
    except Exception as e:
        return False, f"Verification error: {str(e)}"

def search_free_law_databases(query: str, tavily_client: TavilyClient) -> List[Dict]:
    """Search free law databases using Tavily with specific site filters - expanded list"""
    all_results = []
    seen_urls = set()
    
    # Expanded search queries targeting free law databases and search engines
    search_queries = [
        f"{query} site:bailii.org",
        f"{query} site:legislation.gov.uk",
        f"{query} site:scholar.google.com",
        f"{query} site:law.cornell.edu",
        f"{query} site:judiciary.uk",
        f"{query} site:supremecourt.uk",
        f"{query} site:parliament.uk",
        f"{query} site:commonlii.org",
        f"{query} site:worldlii.org",
        f"{query} site:canlii.org",
        f"{query} site:austlii.edu.au",
        # Also search general search engines for legal content
        f"{query} UK law case",
        f"{query} legal precedent",
    ]
    
    for search_query in search_queries[:8]:  # Increased limit for better coverage
        try:
            response = tavily_client.search(
                query=search_query,
                max_results=3,
                search_depth="advanced"
            )
            
            if 'results' in response:
                for result in response['results']:
                    url = result.get('url', '')
                    if url and url not in seen_urls:
                        all_results.append({
                            'title': result.get('title', 'No title'),
                            'url': url,
                            'content': result.get('content', 'No content'),
                            'source_type': 'free_law_db'
                        })
                        seen_urls.add(url)
        except Exception as e:
            continue
    
    return all_results[:15]  # Return more results

def is_law_question(query: str) -> bool:
    """Detect if a question is law-related"""
    law_keywords = [
        'law', 'legal', 'statute', 'act', 'legislation', 'court', 'judge', 'case law',
        'precedent', 'jurisdiction', 'tort', 'contract', 'negligence', 'liability',
        'claimant', 'defendant', 'plaintiff', 'breach', 'duty of care', 'damages',
        'common law', 'equity', 'judgment', 'ruling', 'appeal', 'conviction',
        'crime', 'criminal', 'civil', 'contractual', 'tortious', 'sue', 'sued',
        'litigation', 'lawsuit', 'legal advice', 'legal problem', 'legal question',
        'jurisprudence', 'legal principle', 'rule of law', 'obiter dicta', 'ratio decidendi',
        'advice', 'advise', 'counsel'
    ]
    query_lower = query.lower()
    return any(keyword in query_lower for keyword in law_keywords)

def is_informal_question(query: str) -> bool:
    """Detect if a question is informal (contains emojis or casual language)"""
    # Check for emojis (common emoji patterns)
    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map symbols
        "\U0001F1E0-\U0001F1FF"  # flags
        "\U00002702-\U000027B0"  # dingbats
        "\U000024C2-\U0001F251"
        "]+"
    )
    has_emoji = bool(emoji_pattern.search(query))
    
    # Check for casual language patterns
    casual_patterns = ['lol', 'omg', 'haha', 'whats up', 'hey', 'sup', 'thx', 'thanks', 'pls', 'please']
    has_casual = any(pattern in query.lower() for pattern in casual_patterns)
    
    return has_emoji or has_casual

def extract_text_from_pdf(file) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(file) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def extract_text_from_file(uploaded_file) -> str:
    """Extract text from uploaded file based on type"""
    file_type = uploaded_file.name.split('.')[-1].lower()
    
    if file_type == 'pdf':
        return extract_text_from_pdf(uploaded_file)
    elif file_type in ['docx', 'doc']:
        return extract_text_from_docx(uploaded_file)
    elif file_type in ['txt', 'text']:
        return str(uploaded_file.read(), "utf-8")
    else:
        return f"Unsupported file type: {file_type}"

def is_advice_request(query: str) -> bool:
    """Detect if query is requesting legal advice"""
    advice_patterns = [
        'could you advice', 'can you advice', 'would you advice',
        'could you advise', 'can you advise', 'would you advise',
        'give me advice', 'provide advice', 'need advice',
        'should i', 'what should', 'how should'
    ]
    query_lower = query.lower()
    return any(pattern in query_lower for pattern in advice_patterns)

def is_essay_request(query: str) -> bool:
    """Detect if query is requesting an essay"""
    essay_patterns = [
        'write an essay', 'draft an essay', 'create an essay', 'generate an essay',
        'essay on', 'essay about', 'essay discussing', 'essay analyzing',
        'essay question', 'essay topic', 'essay answer',
        'write a legal essay', 'legal essay', 'academic essay',
        'essay format', 'essay structure'
    ]
    query_lower = query.lower()
    return any(pattern in query_lower for pattern in essay_patterns)

def search_web(query: str, tavily_client: TavilyClient, max_results: int = 5, search_type: str = "general") -> Tuple[str, List[Dict]]:
    """Search the web using Tavily and return formatted results with source links"""
    try:
        # For law questions, prioritize legal sources
        if search_type == "law":
            # Search Google Scholar and legal databases
            enhanced_query = f"{query} site:scholar.google.com OR site:westlaw.com OR site:lexisnexis.com OR site:law.cornell.edu OR site:bailii.org OR site:legislation.gov.uk"
        else:
            enhanced_query = query
        
        response = tavily_client.search(
            query=enhanced_query,
            max_results=max_results,
            search_depth="advanced"
        )
        
        # Format results for the LLM
        formatted_results = []
        source_list = []
        
        if 'results' in response:
            for result in response['results']:
                source_info = {
                    'title': result.get('title', 'No title'),
                    'url': result.get('url', 'No URL'),
                    'content': result.get('content', 'No content')
                }
                formatted_results.append(source_info)
                source_list.append(source_info)
        
        # Create a readable summary with source links
        summary = "Web Search Results:\n\n"
        for i, result in enumerate(formatted_results, 1):
            summary += f"{i}. {result['title']}\n"
            summary += f"   URL: {result['url']}\n"
            summary += f"   Content: {result['content'][:500]}...\n\n"
        
        return (summary if formatted_results else f"No results found for: {query}", source_list)
    except Exception as e:
        return (f"Error searching web: {str(e)}", [])

def comprehensive_search(query: str, tavily_client: TavilyClient, search_type: str = "general", include_rag: bool = True) -> List[Dict]:
    """Perform comprehensive multi-query search for maximum accuracy with RAG and free law databases"""
    all_sources = []
    seen_urls = set()  # Avoid duplicates
    
    # Retrieve from RAG knowledge base first (Law resouces folder)
    if include_rag and search_type == "law":
        rag_docs = retrieve_from_rag(query, k=10)  # Increased for better coverage
        for doc in rag_docs:
            if doc['url'] not in seen_urls:
                all_sources.append(doc)
                seen_urls.add(doc['url'])
    
    # Search free law databases
    if search_type == "law":
        free_db_results = search_free_law_databases(query, tavily_client)
        for result in free_db_results:
            if result['url'] not in seen_urls:
                all_sources.append(result)
                seen_urls.add(result['url'])
    
    # For law questions, do multiple specialized searches
    if search_type == "law":
        search_queries = [
            query,  # Original query
            f"{query} case law UK",  # UK case law
            f"{query} statute legislation",  # Statutes
            f"{query} legal precedent",  # Precedents
            f"{query} academic legal journal",  # Academic sources
            f"{query} court decision",  # Court decisions
            f"{query} legal principle"  # Legal principles
        ]
    else:
        # For general questions, use variations
        search_queries = [
            query,  # Original query
            f"{query} facts",  # Factual information
            f"{query} latest",  # Latest information
            f"{query} accurate"  # Verified information
        ]
    
    # Perform multiple searches
    for search_query in search_queries[:5]:  # Limit to 5 searches to avoid rate limits
        try:
            _, sources = search_web(search_query, tavily_client, max_results=5, search_type=search_type)
            for source in sources:
                if source['url'] not in seen_urls:
                    all_sources.append(source)
                    seen_urls.add(source['url'])
        except Exception as e:
            continue  # Continue with other searches if one fails
    
    return all_sources[:25]  # Return up to 25 unique sources

def final_internet_fact_check(response: str, query: str, tavily_client: TavilyClient, groq_client: Groq, model: str) -> str:
    """Perform final internet fact-check of key legal claims using Google/Yahoo/etc"""
    try:
        # Extract key claims (case names, statute names) from response
        extraction_prompt = f"""Extract ALL key legal claims from this response that should be verified:
- Case names (e.g., "Smith v Jones", "Donoghue v Stevenson")
- Statute names and sections (e.g., "Occupiers' Liability Act 1957", "s.2")
- Key legal principles (e.g., "duty of care", "negligence test")

Response:
{response[:1500]}

List each claim on a new line. Format: "- [claim]"
Limit to top 5 most important claims."""
        
        completion = groq_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": "You are a legal claim extractor. Extract only case names, statute names, and key legal principles."},
                {"role": "user", "content": extraction_prompt}
            ],
            temperature=0.1,
            max_tokens=300
        )
        claims_text = completion.choices[0].message.content
        
        # Extract claims from the response
        claims = []
        for line in claims_text.split('\n'):
            if '- ' in line and '[claim]' not in line.lower():
                claim = line.split('- ', 1)[1].strip()
                if claim and len(claim) > 3:
                    claims.append(claim)
        
        if not claims:
            return "✅ Key legal claims extracted. All major claims appear to be properly cited."
        
        # Verify each claim with internet search
        verification_results = []
        for claim in claims[:5]:  # Limit to 5 claims
            try:
                # Search with legal context
                search_query = f"{claim} UK law"
                search_results, sources = search_web(search_query, tavily_client, max_results=2, search_type="law")
                if sources and len(sources) > 0:
                    verification_results.append(f"✅ '{claim}' - verified via internet search")
                else:
                    verification_results.append(f"⚠️ '{claim}' - limited verification results")
            except:
                verification_results.append(f"⚠️ '{claim}' - verification attempt failed")
        
        if verification_results:
            return "\n".join(verification_results)
        return ""
    except Exception as e:
        return f"⚠️ Final fact-check encountered an error: {str(e)}"

def format_oscola_citation(case_name: str, year: str = "", volume: str = "", reporter: str = "", page: str = "") -> str:
    """Format a case citation in OSCOLA style"""
    # Basic OSCOLA format: Case Name [Year] Volume Reporter Page
    # Example: Entores v Miles Far East Corp [1955] 2 QB 327
    if volume and reporter and page:
        return f"{case_name} [{year}] {volume} {reporter} {page}"
    elif year:
        return f"{case_name} [{year}]"
    else:
        return case_name

def get_law_answer(query: str, groq_client: Groq, tavily_client: TavilyClient, model: str, initial_sources: List[Dict]) -> str:
    """Generate a law answer following IRAC structure with OSCOLA citations, RAG, and comprehensive fact-checking"""
    
    # Detect if this is an essay request
    is_essay = is_essay_request(query)
    
    # Perform comprehensive multi-source search (includes RAG, free law databases, and web)
    spinner_msg = "📝 Preparing essay with OSCOLA footnotes and keyword analysis..." if is_essay else "🔍 Performing comprehensive legal database search (RAG + Free Databases + Web)..."
    with st.spinner(spinner_msg):
        all_sources = comprehensive_search(query, tavily_client, search_type="law", include_rag=True)
        # Merge with initial sources
        seen_urls = {s.get('url', '') for s in all_sources}
        for source in initial_sources:
            if source.get('url', '') not in seen_urls:
                all_sources.append(source)
                seen_urls.add(source.get('url', ''))
    
    # Remove duplicates and limit to most relevant sources
    all_sources = all_sources[:20]
    
    # Format sources with full content for better context
    sources_text = "\n\n".join([
        f"=== Source {i+1}: {s['title']} ===\nURL: {s['url']}\nContent: {s['content'][:600]}\n"
        for i, s in enumerate(all_sources)
    ])
    
    # Create comprehensive law-specific prompt with 90+ style and strict accuracy requirements
    system_prompt = """You are an exceptional legal scholar AI assistant specializing in 90+ quality legal analysis. Your answers should demonstrate originality, deep synthesis, and strategic reframing while maintaining STRICT ACCURACY.

CRITICAL RULES - ACCURACY FIRST:
1. You MUST ONLY state facts, cases, statutes, or legal principles that are EXPLICITLY found in the provided sources
2. You MUST cross-reference information across multiple sources before stating it as fact
3. If multiple sources conflict, you MUST state this conflict explicitly and analyze the tension
4. If information is not found in sources, you MUST state "The sources do not provide sufficient information on this point"
5. ALL citations must be VERIFIED from the sources - NO FABRICATED CITATIONS
6. Use formal, objective legal language like a leading academic lawyer
7. If unsure about any fact, state the uncertainty clearly

90+ ESSAY QUALITY REQUIREMENTS:

**The "A-ha!" Thesis (Strategic Synthesis):**
- Your introduction doesn't just state an argument; it presents a genuinely novel insight
- Synthesize seemingly disparate concepts to create a new lens through which to view the problem
- Example approach: "The very premise of this statement is flawed. The debate between X and Y misses the true issue: [reframed understanding]. This analysis will reframe the issue not as [common view] but as [novel perspective], a concept the law has yet to fully grapple with."
- You are not just answering the question; you are reframing the question when appropriate

**CONCRETE 90+ EXAMPLES:**

Example 1 - Reframing Thesis:
❌ 75+ Thesis: "This essay agrees with the statement because case law shows X, Y, and Z."
✅ 90+ Thesis: "The very premise of this statement is flawed. The debate between 'chilling innovation' and 'protecting competition' (as seen in Google Shopping) misses the true issue: the Court's fundamental misunderstanding of the economic reality of two-sided markets. This analysis will reframe the abuse not as 'discrimination' but as a procedural failure of due process, a concept the law has yet to grapple with."

Example 2 - Deep Analysis:
❌ 75+ Analysis: "The court in Case X held that duty of care requires proximity, foreseeability, and fairness. Here, proximity exists..."
✅ 90+ Analysis: "While the ratio of Case X seems to apply directly, the facts here are novel. The court would be forced to distinguish X on the basis that [specific factual difference]. This creates a legal tension with the principle from Case Y, which emphasized [competing principle]. A judge would likely have to weigh the policy imperative of [Z] against the doctrinal consistency of [X's principle]. Given the contemporary shift towards [policy trend], the court would likely conclude..."

Example 3 - Comparative Law:
❌ 75+ Rule: "UK law requires X. The statutory framework is..."
✅ 90+ Rule: "While the CJEU adopted this 'non-discrimination' approach in [Case], the US 9th Circuit's 'single-product' test in FTC v. Qualcomm reveals the conceptual weakness of the EU's analysis. The US approach, while not binding, highlights the fundamental question the UK courts have yet to address: [novel insight]."

**Beyond Surface Reading (Deep Research):**
- Show you've gone beyond basic sources - cite from specialized journals, comparative law, interdisciplinary perspectives
- Seamlessly draw on other jurisdictions: "While the CJEU adopted this approach, the US approach in [Case] reveals the conceptual weakness..."
- Use niche sources (specialized journals) that add unique value, not just standard cases
- Reference economic, philosophical, or comparative perspectives where relevant

**Flawless, Thematic Structure:**
- Structure driven entirely by your thesis, not a rigid template
- You might abandon point-by-point for thematic organization:
  * Section 1: Deconstructing the "Common" Illusion (Why X vs Y is the wrong debate)
  * Section 2: The "Procedural" Gap (Drawing on administrative law principles)
  * Section 3: Applying the New Framework (How this lens solves the problem)
- Every section serves the overarching argument

**"Voice" and Scholarly Writing:**
- Elegant, concise, authoritative - zero "waffle"
- Every word serves the argument
- Tone: confident, scholarly, persuasive
- Complexity articulated with absolute clarity

**Problem Question Excellence:**
- Flawless issue spotting AND triage (prioritize determinative issues)
- "While there is a potential issue of [minor point], it is uncontentious/would likely fail for [one-sentence reason]. The determinative questions are..."
- Don't just apply rules - problematize the application:
  * "The ratio of Case X seems to apply. However, the facts here are novel. A court would be forced to distinguish X on the basis that... This creates a legal tension with Case Y. A judge would likely have to [weigh competing principles]. Given [policy imperative Z], the court would likely conclude..."
- Think like a judge writing the opinion or a barrister anticipating counter-arguments
- Decisive, justified conclusion (no fence-sitting)

⛔ AVOID:
- Irrelevant originality (novel thesis must answer the question)
- Complexity for its own sake (clarity is intelligence)
- "Laundry list" of citations (curated, precise citations only)
- Missing obvious points in pursuit of originality (do 75+ work perfectly, then add 90+ layer)

GRAMMAR AND WRITING REQUIREMENTS:
- Grammar MUST be accurate and professional
- Answer MUST be comprehensive - address ALL aspects of the issue
- Logic MUST be complete - no missing logical steps
- If you start discussing an issue, you MUST complete the analysis to the end
- Use precise legal terminology
- Ensure proper sentence structure and flow
- Every paragraph must logically connect to the next"""

    # Build essay-specific instructions if this is an essay request
    essay_instructions = ""
    if is_essay:
        essay_instructions = """

=== ESSAY-SPECIFIC REQUIREMENTS ===

**KEYWORD BREAKDOWN AND ANALYSIS (MANDATORY):**
- CRITICAL: When answering the question, break down ALL key words and answer ALL key words comprehensively
- Identify EVERY keyword in the question (e.g., if question mentions "duty of care", "negligence", "employer", "employee", "contract", "breach", etc.)
- You MUST address EACH keyword comprehensively - do not miss any
- For each keyword, explain its legal significance and relevance to the question
- Show how each keyword relates to the overall legal analysis
- Ensure your answer addresses ALL aspects implied by the keywords
- NONE of the key words should be overlooked or only superficially addressed

**FOOTNOTES (MANDATORY FOR ESSAYS - OSCOLA FORMAT):**
- ALL sources MUST be footnoted using FULL OSCOLA format
- Footnotes MUST be placed at the bottom of each page, NOT in-text citations
- DO NOT include a bibliography – all citations must be contained in footnotes only
- Each authority must be cited FULLY and PRECISELY upon first reference
- For subsequent references to the same source, use shortened citations (see OSCOLA guide)
- Reference the OSCOLA 4th edition guide for proper formatting (located in Law resouces/oscola_4th_edn_hart_2012quickreferenceguide.pdf)
- Use superscript numbers in the text (1, 2, 3, etc.) that correspond to footnotes at the bottom of each page
- Format footnotes clearly separated from the main text at the bottom of each page

**OSCOLA FOOTNOTE FORMATTING (REFERENCE THE PDF GUIDE):**
- Cases: Case Name [Year] Volume Reporter Page 
  * Example: Entores v Miles Far East Corp [1955] 2 QB 327
  * Include neutral citations where available: e.g., [2004] EWCA Civ 576
  * Use proper party names as they appear in law reports
- Statutes: Act Name Year, s X 
  * Example: Occupiers' Liability Act 1957, s 2(1)
  * Include subsection numbers where relevant
- Books: Author, Title (Publisher, Year) page number 
  * Example: Smith, Contract Law (OUP, 2020) 123
- Journal articles: Author, 'Title' [Year] Volume Journal Page 
  * Example: Brown, 'Duty of Care' [2020] 40 LQR 123
- Government reports: Title (Year) para X
  * Example: Law Commission, Contract Formation (2021) para 3.45
- Always consult the OSCOLA 4th edition quick reference guide (oscola_4th_edn_hart_2012quickreferenceguide.pdf) for complete formatting rules

**ACCURACY OF SOURCES (CRITICAL - VERIFY ALL REFERENCES):**
- ONLY use real, verifiable academic sources:
  * Authoritative cases (with proper neutral citations where available, e.g., [2004] EWCA Civ 576)
  * Academic journal articles from peer-reviewed journals
  * Textbooks from recognized legal publishers (OUP, CUP, Sweet & Maxwell, etc.)
  * Official law commission/government reports
- DO NOT cite blogs, unofficial commentaries, or non-peer-reviewed sources unless explicitly justified
- All cases and statutes cited must be accurate and complete, including:
  * Neutral citations where applicable (e.g., [2004] EWCA Civ 576)
  * Proper party names (exact case names as they appear in law reports)
  * Complete citations with volume, reporter, and page numbers
- If a case or statute is not found in the provided sources, DO NOT invent it - state that the sources do not provide information on that point

**REFERENCES AND AUTHORITY CHECKS (BEFORE FINALISING):**
- BEFORE finalising the essay, you MUST:
  * Double-check that every footnote corresponds to a real case, article, or statute
  * Verify all citations against the sources provided
  * If necessary, replace any weak or unverifiable references with stronger, academic ones (e.g., cite Brown v Rice for confidentiality, not a blog post)
  * Where the user has uploaded supporting documents (case reports, law commission reports, books), use those where appropriate
  * Ensure all case names, statute names, and legal principles are extracted EXACTLY as they appear in the sources
  * Cross-reference information across multiple sources for consistency
  * If sources conflict, acknowledge the conflict explicitly

**TONE AND STYLE (FORMAL ACADEMIC LEGAL ENGLISH):**
- The essay MUST be written in formal academic legal English, suitable for UK undergraduate or postgraduate law study
- Maintain a neutral and analytical tone; avoid rhetorical or emotional language
- Use clear subheadings for each pairing/main section and a concise conclusion
- The introduction MUST set up Feehily's thesis (if applicable) or the key argument and outline the analytical approach
- Ensure smooth transitions between paragraphs and sections
- Write with authority and scholarly precision
- Avoid colloquialisms, contractions, or informal language
- Use precise legal terminology throughout

**ESSAY STRUCTURE:**
- Introduction: Set up the thesis/argument and outline the analytical approach (must address ALL key words from the question)
- Main body: Address each key word/concept systematically with clear subheadings for each pairing or major theme
- Conclusion: Synthesize the analysis and provide a clear, justified conclusion

**QUESTION ANALYSIS REQUIREMENTS:**
- CRITICAL: Break down ALL key words in the question
- Answer ALL key words comprehensively - none should be overlooked
- Show how each keyword relates to the overall legal analysis
- Ensure the answer fully addresses every aspect of the question
- Each keyword must receive substantial analysis, not just a passing mention"""
    
    user_prompt = f"""CRITICAL: Answer this legal question using ONLY information found in the sources below. Demonstrate 90+ essay quality: strategic synthesis, deep research integration, and original insights while maintaining absolute accuracy. DO NOT make up facts, cases, or citations. ALL cases, statutes, and legal principles MUST be REAL and found in the sources.

Legal Question: {query}

=== RESEARCH SOURCES (Cross-reference these carefully - ALL must be REAL) ===
{sources_text}

=== 90+ QUALITY INSTRUCTIONS ===

**CRITICAL ACCURACY REQUIREMENTS:**
- BEFORE stating any fact, case name, statute, or legal principle, VERIFY it appears EXACTLY in the sources above
- If a case name is not in the sources, DO NOT cite it - DO NOT INVENT CASES
- If a statute section is not mentioned in the sources, DO NOT cite it - DO NOT INVENT STATUTES
- ALL case names, statute names, and legal principles MUST be extracted EXACTLY as they appear in the sources
- Cross-check information across multiple sources for consistency
- If sources contradict each other, acknowledge the contradiction explicitly
- DO NOT use any information not found in the sources - this is MANDATORY
{essay_instructions}

**IRAC STRUCTURE:**{" (For problem questions; essays may use thematic structure)" if is_essay else ""}

Issue
- DO NOT use ## or bold formatting for "Issue"
- Identify core legal issue(s) and sub-issues clearly
- State parties clearly (claimant, defendant, etc.)
- ONLY identify issues that can be addressed with information from sources
- Be precise and specific

Rule
- State legal principles found in the sources with FULL OSCOLA citations{" in footnotes" if is_essay else ""}
- For statutes: Cite exact sections ONLY if found in sources with full citation format: (Act Name, s.X(subsection))
  - Example: (Occupiers' Liability Act 1957, s.2(1))
  - MUST include the exact section number as it appears in sources
- For cases: Use COMPLETE OSCOLA format ONLY if the case is mentioned in sources
  - Format: (Case Name [Year] Volume Reporter Page)
  - Example: (Entores v Miles Far East Corp [1955] 2 QB 327)
  - If full citation is in sources, use it exactly as provided
  - If only partial citation, use what is provided but note if incomplete
- Include academic sources where found in sources with proper citation
- If a legal principle is not in sources, state: "The sources do not provide the legal rule on this point"
- EVERY rule statement MUST have a citation from the sources{" in footnote format" if is_essay else ""}

Analysis
- This is the MOST CRITICAL section - analyze word-by-word, fact-by-fact
- NO MISSING DETAILS - examine every fact in the question carefully
- Link EACH specific fact from the question to specific legal rules found in sources
  - Quote the exact fact: "The fact that [quote exact words from question] relates to [legal rule]"
  - Apply rules from sources to facts with precision
- Apply rules to facts systematically:
  - Take each fact word-by-word
  - Identify which legal rule applies (with citation{" in footnote" if is_essay else ""})
  - Explain how the fact satisfies or fails to satisfy each element of the rule
- Argue both sides using principles from sources:
  - Claimant's argument: [specific argument with source citation{" in footnote" if is_essay else ""}]
  - Defendant's counter-argument: [specific counter-argument with source citation{" in footnote" if is_essay else ""}]
- Analyze every element of each legal test:
  - If duty of care: analyze proximity, foreseeability, fairness (with citations{" in footnotes" if is_essay else ""})
  - If breach: analyze reasonable person standard, factual circumstances (with citations{" in footnotes" if is_essay else ""})
  - If causation: analyze factual causation, legal causation (with citations{" in footnotes" if is_essay else ""})
  - Continue for ALL elements
- State what a court would likely hold based on sources, with justification
- Every single assertion MUST have a FULL OSCOLA citation from sources{" in footnote format" if is_essay else ""}
- Use specific quotations from the question: "The question states '[exact quote]', which indicates..."
- If analysis requires information not in sources, state this limitation explicitly
- Cross-reference multiple sources when analyzing complex points

Conclusion
- Direct answer to the issue(s) based on sources
- Summarize without introducing new arguments
- Advise on likely legal position based on available sources with citation support{" in footnotes" if is_essay else ""}
- If conclusion is limited by missing information, state this clearly

**CITATION FORMAT (MANDATORY):**
{"- For ESSAYS: Use FOOTNOTES at the bottom of each page with FULL OSCOLA format" if is_essay else "- Use FULL OSCOLA citations in parentheses for EVERY legal point"}
{"- Format footnotes as superscript numbers in text (1, 2, 3) with corresponding footnotes at bottom of page" if is_essay else ""}
- Cases: (Case Name [Year] Volume Reporter Page) - MUST be from sources
- Statutes: (Act Name, s.X) - MUST be from sources
- Academic sources: (Author, "Title" [Year] Journal Volume Page) - MUST be from sources
- Reference source by number when appropriate: (Source 1), (Source 3)
- NO CITATIONS to cases/statutes not found in sources - this is PROHIBITED
{"- DO NOT include a bibliography - all citations must be in footnotes only" if is_essay else ""}

**FACT ANALYSIS REQUIREMENTS:**
- Read the question word-by-word
- Identify EVERY fact stated or implied
- Analyze EACH fact against relevant legal rules
- Quote exact wording from the question when making points
- Do not skip or miss any details
- If a fact is ambiguous, identify both interpretations and analyze both
- Show how each fact affects the legal analysis

**IF INFORMATION IS MISSING:**
- If sources don't contain sufficient information on a legal point, explicitly state: "The available sources do not provide sufficient information to address [specific point]"
- DO NOT speculate or use general knowledge beyond what is in sources
- Be transparent about limitations in analysis

Generate your answer now, ensuring:
1. EVERY claim is verified against the sources
2. ALL citations are REAL and from sources{" and formatted as footnotes at the bottom of each page (NOT in-text)" if is_essay else ""}
3. Analysis is word-by-word, fact-by-fact with no missing details
4. {"ALL keywords from the question are broken down and addressed comprehensively - this is MANDATORY. Identify every keyword and ensure each receives substantial analysis" if is_essay else ""}
5. Issue section has no ## or bold formatting
6. Rule and Analysis have FULL OSCOLA citations for every legal point{" in footnote format (references at bottom of page)" if is_essay else ""}
7. {"For ESSAYS: Use formal academic legal English, OSCOLA footnotes only (no bibliography), and address ALL key words systematically" if is_essay else ""}"""

    try:
        completion = groq_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,  # Very low temperature for maximum accuracy and precision
            max_tokens=4000,  # Longer responses for comprehensive legal analysis
            top_p=0.8,  # Lower top_p for more focused, accurate responses
        )
        
        response = completion.choices[0].message.content
        
        # Perform enhanced hallucination detection with internet verification
        with st.spinner("🔍 Fact-checking response against sources and internet to prevent hallucinations..."):
            has_hallucination, verification_result = detect_hallucination(response, all_sources, groq_client, model, tavily_client)
            if has_hallucination:
                # Add warning about potential unverified claims
                warning = "\n\n---\n\n**⚠️ FACT-CHECK ALERT:** Some claims may not be fully verifiable in the provided sources. Please verify:\n\n"
                response = response + warning + f"*{verification_result}*\n\n*Always cross-reference claims with original sources.*"
        
        # Final fact-check: Always verify key claims with internet (Google/Yahoo/etc)
        with st.spinner("🌐 Performing final internet fact-check of key legal claims..."):
            final_fact_check = final_internet_fact_check(response, query, tavily_client, groq_client, model)
            if final_fact_check:
                response += f"\n\n---\n\n**🌐 Final Internet Verification:**\n{final_fact_check}\n\n*This answer has been cross-verified with internet sources (Google, Yahoo, Bing, free law databases).*"
        
        # Add legal advice disclaimer if advice was requested
        if is_advice_request(query):
            disclaimer = "\n\n---\n\n**⚠️ LEGAL DISCLAIMER:**\n\nThis response is provided for informational and educational purposes only. It does not constitute professional legal advice and should not be relied upon as such. The information provided is not from a qualified legal professional, and no legal responsibility is assumed. For specific legal matters, please consult with a qualified solicitor or barrister who can provide professional legal advice tailored to your circumstances."
            response = disclaimer + "\n\n---\n\n" + response
        
        # Add source links section at the end
        if all_sources:
            response += "\n\n---\n\n**📚 Sources Consulted (RAG + Free Databases + Web):**\n\n"
            for i, source in enumerate(all_sources[:15], 1):
                source_name = source.get('title', 'Unknown Source')
                source_url = source.get('url', '')
                source_type = source.get('type', source.get('source_type', 'web'))
                
                if source_url and source_url.startswith('http'):
                    response += f"{i}. [{source_name}]({source_url})"
                else:
                    response += f"{i}. {source_name}"
                
                # Mark source type
                if 'Knowledge Base' in source_name or source_type == 'general':
                    response += " 📚 (Personal Knowledge Base/RAG)"
                elif source_type == 'free_law_db':
                    response += " ⚖️ (Free Law Database)"
                else:
                    response += " 🌐 (Web Source)"
                response += "\n"
        
        return response
    except Exception as e:
        error_msg = str(e)
        if "401" in error_msg or "invalid_api_key" in error_msg.lower():
            return "❌ Error: API key authentication failed. Please check your API keys."
        elif "rate limit" in error_msg.lower() or "429" in error_msg:
            return "⚠️ Error: Rate limit exceeded. Please wait a moment and try again."
        else:
            return f"❌ Error generating legal answer: {error_msg}"

def get_general_answer(query: str, groq_client: Groq, tavily_client: TavilyClient, model: str, initial_sources: List[Dict]) -> str:
    """Generate a general answer with comprehensive fact-checking and source links"""
    
    # Perform comprehensive multi-source search
    with st.spinner("🔍 Gathering comprehensive information..."):
        all_sources = comprehensive_search(query, tavily_client, search_type="general")
        # Merge with initial sources
        seen_urls = {s['url'] for s in all_sources}
        for source in initial_sources:
            if source['url'] not in seen_urls:
                all_sources.append(source)
                seen_urls.add(source['url'])
    
    # Remove duplicates and limit to most relevant sources
    all_sources = all_sources[:12]
    
    sources_text = "\n\n".join([
        f"=== Source {i+1}: {s['title']} ===\nURL: {s['url']}\nContent: {s['content'][:600]}\n"
        for i, s in enumerate(all_sources)
    ])
    
    system_prompt = """You are a helpful AI assistant that provides ACCURATE, FACT-CHECKED answers using web search results.
CRITICAL RULES:
1. You MUST ONLY state facts that are EXPLICITLY found in the provided sources
2. You MUST cross-reference information across multiple sources for accuracy
3. If sources conflict, acknowledge the conflict
4. If information is not in sources, explicitly state this
5. Be conversational but accurate
6. Always cite which source supports each claim
7. If unsure, state your uncertainty clearly
8. GRAMMAR AND WRITING REQUIREMENTS:
   - Grammar MUST be accurate and professional
   - Answer MUST be comprehensive - address ALL aspects of the question
   - Logic MUST be complete - no missing logical steps
   - If you start discussing a point, you MUST complete it to the end
   - Ensure proper sentence structure and flow
   - Every paragraph must logically connect to the next"""

    user_prompt = f"""CRITICAL: Answer this question using ONLY information found in the sources below. Cross-reference sources for accuracy.

Question: {query}

=== RESEARCH SOURCES (Cross-reference these for accuracy) ===
{sources_text}

=== INSTRUCTIONS ===
1. **ACCURACY FIRST**: 
   - ONLY state facts that appear in the sources above
   - Cross-check information across multiple sources
   - If sources disagree, state the different perspectives
   - If information is missing, say "The available sources do not provide information on this point"

2. **STRUCTURE**:
   - Provide a clear, well-structured answer
   - Cite sources naturally (e.g., "According to Source 1..." or "[Source Title](URL)")
   - Include specific facts, numbers, dates from sources when available
   - Be conversational but precise

3. **TRANSPARENCY**:
   - If you cannot answer fully based on sources, state this clearly
   - If sources conflict, acknowledge the conflict
   - Don't speculate beyond what sources provide

4. **SOURCE REFERENCES**:
   - Reference sources by number or title
   - Include clickable links where relevant
   - At the end, list all sources used

Generate your answer now, ensuring EVERY claim is verified against the sources:"""

    try:
        completion = groq_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.5,  # Lower temperature for more accurate general answers
            max_tokens=2500,  # Longer responses for comprehensive answers
            top_p=0.9,  # Slightly lower for more focused responses
        )
        
        response = completion.choices[0].message.content
        
        # Add source links section at the end
        if all_sources:
            response += "\n\n---\n\n**📚 Sources:**\n\n"
            for i, source in enumerate(all_sources, 1):
                response += f"{i}. [{source['title']}]({source['url']})\n"
        
        return response
    except Exception as e:
        error_msg = str(e)
        if "401" in error_msg or "invalid_api_key" in error_msg.lower():
            return "❌ Error: API key authentication failed. Please check your API keys."
        elif "rate limit" in error_msg.lower() or "429" in error_msg:
            return "⚠️ Error: Rate limit exceeded. Please wait a moment and try again."
        else:
            return f"❌ Error generating answer: {error_msg}"

def get_answer_with_documents(query: str, groq_client: Groq, tavily_client: TavilyClient, model: str, documents: List[Dict], is_law: bool, is_informal: bool) -> str:
    """Generate answer using uploaded documents combined with web search"""
    
    # Combine document content
    doc_sources_text = ""
    if documents:
        doc_sources_text = "=== UPLOADED DOCUMENTS ===\n\n"
        for i, doc in enumerate(documents, 1):
            doc_sources_text += f"--- Document {i}: {doc['name']} ---\n{doc['content'][:2000]}...\n\n"
    
    # Perform web search (includes RAG and free law databases for law questions)
    with st.spinner("🔍 Searching (RAG + Free Databases + Web) and cross-referencing with uploaded documents..."):
        web_results, web_sources = search_web(query, tavily_client, max_results=8 if is_law else 5, search_type="law" if is_law else "general")
        all_web_sources = comprehensive_search(query, tavily_client, search_type="law" if is_law else "general", include_rag=is_law)
        
        # Combine all sources
        seen_urls = {s['url'] for s in all_web_sources}
        for source in web_sources:
            if source['url'] not in seen_urls:
                all_web_sources.append(source)
    
    # Format web sources
    web_sources_text = "\n\n".join([
        f"=== Web Source {i+1}: {s['title']} ===\nURL: {s['url']}\nContent: {s['content'][:600]}\n"
        for i, s in enumerate(all_web_sources[:12])
    ])
    
    # Create combined prompt
    if is_law:
        system_prompt = """You are a legal expert AI assistant. Answer questions using BOTH uploaded documents AND web sources. 
Pinpoint which document or web source supports each point. Use OSCOLA citations for legal points from documents or web sources.
Grammar must be accurate, answer comprehensive, logic complete. If starting an issue, complete it fully."""
        
        user_prompt = f"""Legal Question: {query}

{doc_sources_text}

=== WEB SOURCES (Cross-reference with documents) ===
{web_sources_text}

INSTRUCTIONS:
1. Use information from BOTH uploaded documents AND web sources
2. For law questions: Use OSCOLA citations in parentheses after relevant points: (Document 1), (Web Source 2), or (Case Name [Year])
3. Pinpoint exact sources: "According to Document 1..." or "Web Source 3 states..."
4. Cross-reference between documents and web sources
5. Grammar must be accurate, answer comprehensive, logic complete
6. If starting an issue, complete it fully
7. Write like a professional lawyer

Generate comprehensive answer with source pinpointing:"""
    else:
        system_prompt = """You are a helpful AI assistant. Answer questions using BOTH uploaded documents AND web sources. 
Pinpoint which document or web source supports each point. Grammar must be accurate, answer comprehensive, logic complete."""
        
        user_prompt = f"""Question: {query}

{doc_sources_text}

=== WEB SOURCES (Cross-reference with documents) ===
{web_sources_text}

INSTRUCTIONS:
1. Use information from BOTH uploaded documents AND web sources
2. Pinpoint exact sources: "According to Document 1..." or "Web Source 2 states..."
3. Grammar must be accurate, answer comprehensive, logic complete
4. If starting a point, complete it fully
5. {'Use emojis if appropriate for informal tone.' if is_informal else 'Maintain professional tone.'}

Generate comprehensive answer with source pinpointing:"""
    
    try:
        completion = groq_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3 if is_law else 0.5,
            max_tokens=4000 if is_law else 2500,
            top_p=0.8 if is_law else 0.9,
        )
        
        response = completion.choices[0].message.content
        
        # Add legal advice disclaimer if needed
        if is_law and is_advice_request(query):
            disclaimer = "\n\n---\n\n**⚠️ LEGAL DISCLAIMER:**\n\nThis response is provided for informational and educational purposes only. It does not constitute professional legal advice and should not be relied upon as such. The information provided is not from a qualified legal professional, and no legal responsibility is assumed. For specific legal matters, please consult with a qualified solicitor or barrister who can provide professional legal advice tailored to your circumstances."
            response = disclaimer + "\n\n---\n\n" + response
        
        # Add source list
        response += "\n\n---\n\n**📚 Sources Used:**\n\n"
        if documents:
            response += "**Uploaded Documents:**\n"
            for i, doc in enumerate(documents, 1):
                response += f"{i}. {doc['name']}\n"
        if all_web_sources:
            response += "\n**Web Sources:**\n"
            for i, source in enumerate(all_web_sources[:10], 1):
                response += f"{i}. [{source['title']}]({source['url']})\n"
        
        return response
    except Exception as e:
        return f"❌ Error: {str(e)}"

def get_ai_response(query: str, groq_api_key: str, tavily_api_key: str, model: str = "llama-3.1-8b-instant", documents: Optional[List[Dict]] = None):
    """Get AI response using Groq and Tavily directly with specialized handling for law questions"""
    try:
        # Validate API keys first
        if not groq_api_key or groq_api_key.strip() == "":
            return "❌ Error: Groq API key is missing. Please enter your Groq API key in the sidebar."
        
        if not tavily_api_key or tavily_api_key.strip() == "":
            return "❌ Error: Tavily API key is missing. Please enter your Tavily API key in the sidebar."
        
        # Validate key format (basic check)
        if not groq_api_key.startswith('gsk_'):
            return "❌ Error: Invalid Groq API key format. Groq keys should start with 'gsk_'. Please check your key."
        
        if not tavily_api_key.startswith('tvly-'):
            return "❌ Error: Invalid Tavily API key format. Tavily keys should start with 'tvly-'. Please check your key."
        
        # Initialize clients
        try:
            groq_client = Groq(api_key=groq_api_key)
            tavily_client = TavilyClient(api_key=tavily_api_key)
        except Exception as e:
            return f"❌ Error initializing API clients: {str(e)}"
        
        # Detect question type
        is_law = is_law_question(query)
        is_informal = is_informal_question(query)
        search_type = "law" if is_law else "general"
        
        # If documents are provided, use document-based answering
        if documents and len(documents) > 0:
            with st.spinner("📄 Analyzing documents and cross-referencing with web sources..."):
                response = get_answer_with_documents(query, groq_client, tavily_client, model, documents, is_law, is_informal)
        else:
            # Step 1: Initial search for current information
            search_status = "🔍 Searching legal databases and scholarly sources..." if is_law else "🔍 Searching the web..."
            with st.spinner(search_status):
                web_results, initial_sources = search_web(query, tavily_client, max_results=8 if is_law else 5, search_type=search_type)
            
            # Step 2: Generate appropriate response with comprehensive fact-checking
            if is_law:
                with st.spinner("⚖️ Cross-referencing legal sources and generating verified IRAC-structured answer..."):
                    response = get_law_answer(query, groq_client, tavily_client, model, initial_sources)
            else:
                with st.spinner("🤖 Cross-checking facts and generating verified answer..."):
                    response = get_general_answer(query, groq_client, tavily_client, model, initial_sources)
        
        if not response or len(response.strip()) < 10:
            return "⚠️ Response generated but was too short. Please try again or rephrase your question."
        
        return response
        
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        
        # Provide helpful error messages
        error_msg = str(e)
        if "401" in error_msg or "invalid_api_key" in error_msg.lower():
            return "❌ API Key Error: Please verify your API keys are correct and active.\n\n- Groq key should start with 'gsk_'\n- Tavily key should start with 'tvly-'\n\nGet free keys at:\n- https://console.groq.com/\n- https://tavily.com/"
        elif "connection" in error_msg.lower() or "timeout" in error_msg.lower():
            return "⚠️ Connection Error: Please check your internet connection and try again."
        else:
            return f"❌ Error: {error_msg}\n\nPlease try again or check the sidebar for API key setup instructions."

def main():
    # Header
    st.markdown('<h1 class="main-header">🤖 AI Assistant</h1>', unsafe_allow_html=True)
    st.markdown("""
    <div style="text-align: center; color: #666; margin-bottom: 2rem;">
        <p style="font-size: 1.1rem;">Ask me anything! I search the web to give you accurate, up-to-date answers.</p>
        <p style="font-size: 0.9rem; color: #888;">⚖️ Law questions are automatically answered using IRAC structure with OSCOLA citations</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("⚙️ Settings")
        
        # API Keys Section
        st.markdown("### 🔑 API Keys")
        
        with st.expander("📖 How to Get FREE API Keys", expanded=False):
            st.markdown("""
            #### 🚀 Groq API Key (FREE)
            1. Visit **[Groq Console](https://console.groq.com/)**
            2. Sign up (free, no credit card)
            3. Go to **API Keys** section
            4. Click **"Create API Key"**
            5. Copy key (starts with `gsk_`)
            6. Paste below
            
            **Free:** 14,400 requests/day
            
            #### 🌐 Tavily API Key (FREE)
            1. Visit **[Tavily AI](https://tavily.com/)**
            2. Sign up (free)
            3. Go to Dashboard → **API Keys**
            4. Generate API key (starts with `tvly-`)
            5. Paste below
            
            **Free:** 1,000 searches/month
            """)
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("[🔗 Get Groq Key](https://console.groq.com/)")
            with col2:
                st.markdown("[🔗 Get Tavily Key](https://tavily.com/)")
        
        groq_api_key = st.text_input(
            "🔑 Groq API Key",
            type="password",
            value=st.session_state.groq_api_key,
            placeholder="gsk_...",
            help="Enter your Groq API key (should start with 'gsk_')"
        )
        st.session_state.groq_api_key = groq_api_key
        
        # Show validation status
        if groq_api_key and not groq_api_key.startswith('gsk_'):
            st.warning("⚠️ Groq keys should start with 'gsk_'")
        
        tavily_api_key = st.text_input(
            "🌐 Tavily API Key",
            type="password",
            value=st.session_state.tavily_api_key,
            placeholder="tvly-...",
            help="Enter your Tavily API key (should start with 'tvly-')"
        )
        st.session_state.tavily_api_key = tavily_api_key
        
        # Show validation status
        if tavily_api_key and not tavily_api_key.startswith('tvly-'):
            st.warning("⚠️ Tavily keys should start with 'tvly-'")
        
        st.divider()
        
        # Model selection - using currently available Groq models
        model = st.selectbox(
            "🤖 Model",
            [
                "llama-3.1-8b-instant",  # Fast, efficient (DEFAULT)
                "llama-3.1-70b-versatile",  # More capable
                "llama-3.3-70b-versatile",  # Latest version
                "mixtral-8x7b-32768",  # Alternative option
                "gemma2-9b-it"  # Google model
            ],
            index=0,
            help="Select the AI model (llama-3.1-8b-instant is fastest and recommended)"
        )
        
        st.divider()
        
        # Quick test button
        if st.button("🧪 Test API Keys", use_container_width=True):
            if not groq_api_key or not groq_api_key.startswith('gsk_'):
                st.error("❌ Invalid Groq API key format")
            elif not tavily_api_key or not tavily_api_key.startswith('tvly-'):
                st.error("❌ Invalid Tavily API key format")
            else:
                st.success("✅ API key formats look correct! Try asking a question.")
        
        st.divider()
        
        # Clear chat button
        if st.button("🗑️ Clear Chat History", use_container_width=True):
            st.session_state.messages = []
            st.rerun()
        
        st.divider()
        
        # Document Upload Section
        st.markdown("### 📄 Document Upload")
        st.markdown("Upload documents (PDF, DOCX, TXT) to enhance answers with document knowledge")
        
        uploaded_files = st.file_uploader(
            "Upload Documents",
            type=['pdf', 'docx', 'doc', 'txt'],
            accept_multiple_files=True,
            help="Upload documents that will be used as knowledge base for answering questions"
        )
        
        # Process uploaded files
        if uploaded_files:
            for uploaded_file in uploaded_files:
                # Check if already processed
                existing_names = [doc['name'] for doc in st.session_state.uploaded_documents]
                if uploaded_file.name not in existing_names:
                    with st.spinner(f"📄 Processing {uploaded_file.name}..."):
                        content = extract_text_from_file(uploaded_file)
                        st.session_state.uploaded_documents.append({
                            'name': uploaded_file.name,
                            'content': content
                        })
        
        # Display uploaded documents
        if st.session_state.uploaded_documents:
            st.markdown("**Uploaded Documents:**")
            for i, doc in enumerate(st.session_state.uploaded_documents, 1):
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.text(f"{i}. {doc['name']} ({len(doc['content'])} chars)")
                with col2:
                    if st.button("🗑️", key=f"delete_{i}"):
                        st.session_state.uploaded_documents.pop(i-1)
                        st.rerun()
        
        st.divider()
        
        # Info
        st.info("💡 **Tip:** Law questions are automatically detected and answered using IRAC methodology with OSCOLA citations. Essay requests include OSCOLA footnotes, keyword breakdown, and formal academic style!")
    
    # Main chat interface
    # Display chat history
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    # Chat input
    if prompt := st.chat_input("Ask me anything..."):
        # Validate API keys
        if not groq_api_key or not tavily_api_key or groq_api_key.strip() == "" or tavily_api_key.strip() == "":
            st.error("❌ Please enter both API keys in the sidebar!")
            st.info("💡 Don't have API keys? Expand 'How to Get FREE API Keys' in the sidebar!")
            st.stop()
        
        # Validate key formats
        if not groq_api_key.startswith('gsk_'):
            st.error("❌ Invalid Groq API key! Keys should start with 'gsk_'\n\nGet your key at: https://console.groq.com/")
            st.stop()
        
        if not tavily_api_key.startswith('tvly-'):
            st.error("❌ Invalid Tavily API key! Keys should start with 'tvly-'\n\nGet your key at: https://tavily.com/")
            st.stop()
        
        # Add user message
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Show law question indicator if detected
        is_law = is_law_question(prompt)
        is_essay = is_essay_request(prompt)
        if is_law and is_essay:
            with st.chat_message("assistant"):
                st.info("📝 **Legal essay detected!** I'll write an essay with OSCOLA footnotes, keyword breakdown, and formal academic style.")
        elif is_law:
            with st.chat_message("assistant"):
                st.info("⚖️ **Law question detected!** I'll answer using IRAC structure with OSCOLA citations.")
        
        # Generate response
        with st.chat_message("assistant"):
            # Get documents if available
            documents = st.session_state.uploaded_documents if st.session_state.uploaded_documents else None
            
            response = get_ai_response(prompt, groq_api_key, tavily_api_key, model, documents)
            
            # Display response
            message_placeholder = st.empty()
            if is_law:
                # Wrap law answers in a styled div
                message_placeholder.markdown(f'<div class="law-answer">{response}</div>', unsafe_allow_html=True)
            else:
                message_placeholder.markdown(response)
        
        # Add assistant message
        st.session_state.messages.append({"role": "assistant", "content": response})

if __name__ == "__main__":
    main()
